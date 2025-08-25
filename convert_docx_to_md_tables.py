from pathlib import Path
from docx import Document
import re

INPUT_DIR = Path("Nivel2_Leccion1")
OUTPUT_DIR = Path("markdown_nivel2_leccion1")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

HEADING_MAP = {
    "heading 1": "#",
    "heading 2": "##",
    "heading 3": "###",
    "heading 4": "####"
}

def map_heading(style_name: str):
    if not style_name:
        return None
    ln = style_name.lower()
    for k, sym in HEADING_MAP.items():
        if k in ln or k.replace(" ", "") in ln or ("título" in ln and k.split()[-1] in ln):
            return sym
    return None

def clean_filename(name: str) -> str:
    base = Path(name).stem
    base = re.sub(r"[^\w\-]+", "_", base)
    return base + ".md"

def table_to_markdown(table):
    rows = []
    for r in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in r.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Si primera fila parece encabezado
    header = rows[0]
    md = []
    md.append("| " + " | ".join(header) + " |")
    md.append("| " + " | ".join(["---"] * len(header)) + " |")

    for row in rows[1:]:
        md.append("| " + " | ".join(row) + " |")
    return "\n".join(md) + "\n"

def doc_to_markdown(doc: Document) -> str:
    # Mezclar párrafos y tablas en orden
    md_lines = []
    # Acceso directo al XML: cada elemento del cuerpo es p o tbl
    for block in doc.element.body:
        tag = block.tag.split('}')[-1]
        if tag == "p":
            # Encontrar el índice de ese paragraph en doc.paragraphs
            # Método: recorrer en orden
            # (Más simple: extraer texto usando un índice incremental)
            pass

    # Estrategia más simple: usar doc.paragraphs para texto y luego insertar tablas en orden
    # Para orden correcto necesitamos intercalar; recreamos orden:
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl

    for child in doc.element.body:
        if isinstance(child, CT_P):
            # Convertir a paragraph
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            text = p.text.strip()
            if not text:
                md_lines.append("")
                continue
            heading = map_heading(p.style.name if p.style else "")
            if heading:
                md_lines.append(f"{heading} {text}")
            else:
                md_lines.append(text)
        elif isinstance(child, CT_Tbl):
            from docx.table import _Cell, Table
            table = Table(child, doc)
            md_lines.append(table_to_markdown(table))
            md_lines.append("")  # separación

    # Compactar blancos
    cleaned = []
    prev_blank = False
    for l in md_lines:
        if l.strip() == "":
            if not prev_blank:
                cleaned.append("")
            prev_blank = True
        else:
            cleaned.append(l)
            prev_blank = False

    return "\n".join(cleaned).strip() + "\n"

def convert_all():
    docs = list(INPUT_DIR.glob("*.docx"))
    if not docs:
        print(f"No se encontraron .docx en {INPUT_DIR.resolve()}")
        return
    for docx_file in docs:
        doc = Document(docx_file)
        md_text = doc_to_markdown(doc)
        out_file = OUTPUT_DIR / clean_filename(docx_file.name)
        out_file.write_text(md_text, encoding="utf-8")
        print(f"Convertido (con tablas): {docx_file.name} -> {out_file}")

if __name__ == "__main__":
    convert_all()